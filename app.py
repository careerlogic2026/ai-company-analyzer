import streamlit as st
import google.generativeai as genai
import requests
from bs4 import BeautifulSoup
import PyPDF2
from io import BytesIO
from docx import Document
from tavily import TavilyClient
import time
import re
import datetime
from urllib.parse import urljoin, urlparse

# ==========================================
# 0. 初期設定とUIスタイル
# ==========================================
st.set_page_config(page_title="AI企業分析: 二刀流ディープクローラー", layout="wide")
st.markdown("""
    <style>
    .main { background-color: #f8f9fa; }
    .source-card { background-color: white; padding: 12px; border-radius: 8px; border-left: 5px solid #007bff; margin-bottom: 8px; box-shadow: 0 1px 3px rgba(0,0,0,0.1); }
    .ext-badge { background-color: #ff5722; color: white; padding: 3px 6px; border-radius: 4px; font-size: 0.75em; margin-right: 8px; font-weight: bold; }
    .int-badge { background-color: #007bff; color: white; padding: 3px 6px; border-radius: 4px; font-size: 0.75em; margin-right: 8px; font-weight: bold; }
    .date-badge { background-color: #28a745; color: white; padding: 3px 6px; border-radius: 4px; font-size: 0.75em; margin-right: 8px; font-weight: bold; }
    </style>
    """, unsafe_allow_html=True)

st.sidebar.title("🛠️ APIキー設定")
gemini_key = st.sidebar.text_input("Gemini API Key", value=st.secrets.get("GEMINI_API_KEY", ""), type="password")
tavily_key = st.sidebar.text_input("Tavily API Key", value=st.secrets.get("TAVILY_API_KEY", ""), type="password")

# ステータス管理
if "step" not in st.session_state: st.session_state.step = 1
if "l1_results" not in st.session_state: st.session_state.l1_results = {}
if "ext_results" not in st.session_state: st.session_state.ext_results = []

# ==========================================
# 1. クローラー＆抽出エンジン
# ==========================================

# エンジンA-1: 第1階層（L1）の浅堀りクローラー
def get_l1_links(base_url, max_links=40):
    try:
        res = requests.get(base_url, timeout=10)
        soup = BeautifulSoup(res.text, 'html.parser')
        base_parsed = urlparse(base_url)
        
        links = set()
        for a in soup.find_all('a', href=True):
            href = a['href']
            # 相対パスを絶対URLに変換し、#（アンカー）を除去
            full_url = urljoin(base_url, href).split('#')[0]
            full_parsed = urlparse(full_url)
            
            # 安全装置①：同じドメイン＆「親URLのパス」から始まる下層ページのみ抽出
            if full_parsed.netloc == base_parsed.netloc and full_parsed.path.startswith(base_parsed.path):
                if full_url != base_url: # 自分自身は除く
                    links.add(full_url)
                    
        return list(links)[:max_links]
    except Exception as e:
        return []

# エンジンA-2: ディープクローラー（L2〜L3の深掘りとテキスト抽出）
def deep_crawl_extract(start_url, max_depth=2, max_pages=5):
    """チェックされたディレクトリを起点に、指定階層まで潜ってテキストを根こそぎ回収する"""
    visited = set()
    queue = [(start_url, 1)] # (URL, 現在の深さ)
    all_text = ""
    
    while queue and len(visited) < max_pages:
        current_url, depth = queue.pop(0)
        if current_url in visited: continue
        
        visited.add(current_url) # 安全装置③：無限ループ防止
        
        try:
            res = requests.get(current_url, timeout=10)
            soup = BeautifulSoup(res.text, 'html.parser')
            
            # ノイズ除去とテキスト抽出
            for s in soup(['script', 'style', 'nav', 'header', 'footer']): s.decompose()
            page_text = soup.get_text(separator=' ', strip=True)[:3000] # 長すぎる場合はカット
            all_text += f"\n[URL: {current_url}]\n{page_text}\n"
            
            # まだ深掘りできるなら、そのページ内の子リンクをキューに追加
            if depth < max_depth:
                parsed_curr = urlparse(current_url)
                for a in soup.find_all('a', href=True):
                    full_url = urljoin(current_url, a['href']).split('#')[0]
                    parsed_full = urlparse(full_url)
                    # 同じディレクトリ配下に限定
                    if parsed_full.netloc == parsed_curr.netloc and parsed_full.path.startswith(parsed_curr.path):
                        if full_url not in visited:
                            queue.append((full_url, depth + 1))
        except Exception:
            pass
            
    return all_text

# 単一URL用（外部メディア用）
def extract_single_url(url):
    try:
        headers = {"User-Agent": "Mozilla/5.0"}
        if url.lower().endswith(".pdf"):
            res = requests.get(url, headers=headers, timeout=15)
            reader = PyPDF2.PdfReader(BytesIO(res.content))
            return "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])[:10000]
        res = requests.get(url, headers=headers, timeout=10)
        soup = BeautifulSoup(res.content, 'html.parser')
        for s in soup(['script', 'style', 'nav', 'header', 'footer']): s.decompose()
        return soup.get_text(separator='\n', strip=True)[:5000]
    except Exception: return "[抽出スキップ]"

# ==========================================
# 2. UI フェーズ1: ターゲット指定
# ==========================================
st.title("🎯 企業分析AI: 二刀流ディープクローラー")
st.write("自社サイトの深層（オウンド）と、著名ビジネスメディアの客観的評価（アーンド）を両面からしゃぶり尽くします。")

if st.session_state.step == 1:
    company_name = st.text_input("🏢 企業名 (必須)", value="株式会社マイナビ")
    
    st.markdown("### 📁 A. 【自社サイト】 深掘りする起点URL（複数可）")
    st.info("ここに指定したURLの「1階層下」を一覧化します。（例：ニュース一覧ページ、新卒採用トップページなど）")
    base_urls_input = st.text_area(
        "起点URLを改行で入力してください",
        value="https://www.mynavi.jp/news/\nhttps://www.mynavi.jp/recruit/newgraduate/\n",
        height=100
    )

    if st.button("🚀 ディレクトリ探索 ＆ 外部メディア検索を開始", type="primary"):
        if not gemini_key or not tavily_key: st.error("APIキーを設定してください。"); st.stop()
        
        client = TavilyClient(api_key=tavily_key)
        now = datetime.datetime.now()
        years = f"({now.year} OR {now.year-1})"
        short_name = company_name.replace("株式会社", "").replace("合同会社", "").strip()

        with st.spinner("自社ディレクトリの浅堀りと、著名ビジネスメディアの検索を同時実行中..."):
            
            # --- エンジンA：自社サイトのL1クローリング ---
            base_urls = [u.strip() for u in base_urls_input.split('\n') if u.strip()]
            l1_results = {}
            for b_url in base_urls:
                l1_results[b_url] = get_l1_links(b_url)
            
            # --- エンジンB：外部ビジネスメディアの指名打ち ---
            # 著名メディアのみを厳選
            top_media_domains = [
                "nikkei.com", "xtrend.nikkei.com", "toyokeizai.net", "diamond.jp", 
                "forbesjapan.com", "newspicks.com", "itmedia.co.jp", "businessinsider.jp", 
                "bridge.jp.net", "prtimes.jp"
            ]
            q_ext = f"{short_name} (戦略 OR 提携 OR インタビュー OR 業績 OR 新サービス) {years}"
            
            try:
                ext_hits = client.search(query=q_ext, include_domains=top_media_domains, max_results=50).get("results", [])
                
                # 日付抽出とソート
                for r in ext_hits:
                    d = re.search(r'(202[4-6])[年/.-](1[0-2]|0?[1-9])[月/.-](3[01]|[12][0-9]|0?[1-9])', r.get('content',''))
                    r['date_str'] = d.group(0) if d else f"直近2年"
                ext_hits.sort(key=lambda x: x.get('date_str',''), reverse=True)
                st.session_state.ext_results = ext_hits
            except Exception as e:
                st.session_state.ext_results = []
                st.warning(f"外部検索エラー: {e}")

            st.session_state.company_name = company_name
            st.session_state.l1_results = l1_results
            st.session_state.step = 2
            st.rerun()

# ==========================================
# 3. UI フェーズ2: 取捨選択（アコーディオン）
# ==========================================
elif st.session_state.step == 2:
    st.subheader("📋 分析対象ディレクトリと記事の取捨選択")
    st.write("AIが抽出したリストです。深掘り（ディープクロール）させたいディレクトリや、読ませたい記事にチェックを入れてください。")
    
    selected_internal = []
    selected_external = []

    # 📁 自社サイトのアコーディオン
    st.markdown("#### 📁 【自社サイト】 深掘り候補ディレクトリ")
    for base_url, links in st.session_state.l1_results.items():
        with st.expander(f"🔻 起点: {base_url} (検出: {len(links)}件)", expanded=True):
            if not links: st.write("下層リンクが見つかりませんでした。")
            for j, link in enumerate(links):
                st.markdown(f"<div class='source-card'><span class='int-badge'>内部</span> <small>{link}</small></div>", unsafe_allow_html=True)
                if st.checkbox("このディレクトリをさらに2階層深掘りする", key=f"int_{base_url}_{j}", value=(j<5)):
                    selected_internal.append(link)

    # 📰 外部メディアのアコーディオン
    st.markdown("#### 📰 【外部メディア】 著名ビジネス記事")
    with st.expander(f"🔻 客観的評価・PR記事 (検出: {len(st.session_state.ext_results)}件)", expanded=True):
        if not st.session_state.ext_results: st.write("該当記事が見つかりませんでした。")
        for j, hit in enumerate(st.session_state.ext_results):
            st.markdown(f"<div class='source-card'><span class='ext-badge'>外部</span> <span class='date-badge'>📅 {hit['date_str']}</span> <b>{hit['title']}</b><br><a href='{hit['url']}' target='_blank'><small>{hit['url']}</small></a></div>", unsafe_allow_html=True)
            if st.checkbox("この記事を分析に含める", key=f"ext_{j}", value=(j<10)):
                selected_external.append(hit['url'])

    if st.button("🚀 選択項目でディープクロール＆レポート作成を開始", type="primary"):
        if not selected_internal and not selected_external: 
            st.error("最低1つの項目を選択してください。")
        else:
            st.session_state.selected_internal = selected_internal
            st.session_state.selected_external = selected_external
            st.session_state.step = 3
            st.rerun()

    if st.button("🔙 やり直す"):
        st.session_state.step = 1
        st.rerun()

# ==========================================
# 4. UI フェーズ3 & 4: 深層クロールと分析実行
# ==========================================
elif st.session_state.step == 3:
    st.subheader("🧠 ディープクローリング ＆ 最終分析プロセス")
    genai.configure(api_key=gemini_key)
    model = genai.GenerativeModel('gemini-3.1-flash-lite-preview')
    
    all_extracted_facts = ""

    with st.status("🕵️‍♂️ エージェントがディープクローリングと情報抽出を実行中...", expanded=True) as status:
        
        # 1. 自社サイトの深掘り実行
        for i, url in enumerate(st.session_state.selected_internal):
            st.write(f"📁 階層深掘り中({i+1}/{len(st.session_state.selected_internal)}): {url}")
            # L2〜L3まで潜って最大5ページ分のテキストを結合して取得
            deep_text = deep_crawl_extract(url, max_depth=2, max_pages=5)
            
            if deep_text.strip():
                res = model.generate_content(f"自社データ:\n{deep_text}\n任務:事業構造、戦略、組織に関する事実を具体的に箇条書きで抽出せよ。")
                all_extracted_facts += f"\n--- 【自社深掘り】 {url}配下 ---\n{res.text}\n"

        # 2. 外部メディアの精読実行
        for i, url in enumerate(st.session_state.selected_external):
            st.write(f"📰 外部記事精読中({i+1}/{len(st.session_state.selected_external)}): {url}")
            ext_text = extract_single_url(url)
            
            if ext_text.strip() and "[抽出スキップ]" not in ext_text:
                res = model.generate_content(f"外部記事:\n{ext_text}\n任務:この記事から読み取れる企業の戦略的意図や客観的評価を抽出せよ。")
                all_extracted_facts += f"\n--- 【外部メディア】 {url} ---\n{res.text}\n"
                
        status.update(label="✅ 全情報の収集・精読が完了しました", state="complete")

    # 3. リレー執筆（コンテキスト・チェーン）
    st.subheader("📊 最終統合分析レポート")
    report_md = ""
    context = ""
    
    EDITOR_PROMPTS = {
        "1. 業界構造と事業モデルの深層": "1500文字以上で執筆。自社の発信と外部メディアの評価を統合し、収益エコシステムを解説せよ。",
        "2. 事業進化論と最新戦略動向": "1500文字以上で執筆。沿革と直近のPR・提携事例を繋ぎ、未来への成長物語を論理的に描け。",
        "3. 組織文化とコア人材の生態": "1500文字以上で執筆。採用ページの深掘りデータ等から、現場のリアルな熱量や求める人物像を言語化せよ。",
        "4. キャリア価値と市場評価の推論": "1500文字以上で執筆。30歳時点での想定年収、スカウトされ得る業界、身につく独自スキルを徹底推論せよ。"
    }

    with st.status("👑 編集長AIが圧倒的熱量でレポートを編纂中...", expanded=True) as ed_status:
        for title, prompt in EDITOR_PROMPTS.items():
            st.write(f"🖋️ {title} を執筆中...")
            final_res = model.generate_content(f"【収集した全事実】\n{all_extracted_facts}\n\n【既出の章(重複を避けること)】\n{context}\n\n章タイトル:{title}\n指示:{prompt}\n※プロのコンサルタントとして、事実に基づき重厚に執筆せよ。")
            report_md += f"## {title}\n\n{final_res.text}\n\n---\n\n"
            context += f"【{title}要約】\n{final_res.text[:300]}...\n"
            
        ed_status.update(label="🎉 究極のハイブリッド分析レポートが完成しました！", state="complete")

    # 4. 結果表示とダウンロード
    st.markdown(report_md)
    
    doc = Document(); doc.add_heading(f"{st.session_state.company_name} 二刀流ディープ分析レポート", 0)
    for p in report_md.split('\n'):
        if p.startswith('## '): doc.add_heading(p[3:], 1)
        elif p.strip() and p != "---": doc.add_paragraph(p)
    bio = BytesIO(); doc.save(bio)
    st.download_button("📄 レポート(Word)を保存", data=bio.getvalue(), file_name=f"{st.session_state.company_name}_HybridReport.docx", type="primary")
    
    if st.button("🔄 新しい分析を始める"):
        st.session_state.step = 1
        st.rerun()
